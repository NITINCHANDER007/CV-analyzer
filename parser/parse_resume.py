"""
parse_resume.py
---------------
Parses a PDF or DOCX resume into structured JSON using:
  - pdfplumber  : raw text extraction from PDF
  - python-docx : raw text extraction from DOCX
  - Gemini API  : NLP structuring into predefined schema
"""

import os
import json
import re
import logging
from pathlib import Path
from typing import Optional

import pdfplumber
from docx import Document
import google.generativeai as genai

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s — %(message)s",
)
logger = logging.getLogger("parse_resume")

# ---------------------------------------------------------------------------
# Gemini configuration
# ---------------------------------------------------------------------------
_GEMINI_MODEL = "gemini-3.6-flash"

EXTRACTION_PROMPT = """
You are an expert resume parser. Given the raw text of a resume below, extract
the information and return ONLY a valid JSON object with exactly these fields:

{{
  "name": "<full name as a string>",
  "email": "<email address as a string or null>",
  "phone": "<phone number as a string or null>",
  "skills": ["<skill1>", "<skill2>", "..."],
  "experience": [
    {{
      "role": "<job title>",
      "company": "<company name>",
      "years": "<duration, e.g. '2020-2023' or '2 years'>"
    }}
  ],
  "projects": [
    {{
      "title": "<project name>",
      "description": "<one or two sentence description>"
    }}
  ],
  "education": [
    {{
      "degree": "<degree / qualification>",
      "institution": "<university or school name>",
      "year": "<graduation year or date range>"
    }}
  ]
}}

Rules:
- Return ONLY the JSON object, with no markdown fences, no commentary, no extra text.
- If a field has no information, use null for strings or [] for arrays.
- Skills must be a flat list of strings (no nested objects).
- Do not invent information not present in the resume.

--- RESUME TEXT START ---
{resume_text}
--- RESUME TEXT END ---
"""

# ---------------------------------------------------------------------------
# Text extraction helpers
# ---------------------------------------------------------------------------


def extract_text_from_pdf(file_path: Path) -> str:
    """Extract all text from a PDF file using pdfplumber."""
    logger.info("Extracting text from PDF: %s", file_path)
    pages_text: list[str] = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                text = page.extract_text()
                if text:
                    pages_text.append(text)
                    logger.debug("  Page %d: %d chars extracted", i, len(text))
                else:
                    logger.debug("  Page %d: no text found (may be image-only)", i)
    except Exception as exc:
        raise RuntimeError(f"Failed to read PDF '{file_path}': {exc}") from exc

    if not pages_text:
        raise ValueError(
            f"No extractable text found in '{file_path}'. "
            "The PDF may be scanned / image-based."
        )
    return "\n".join(pages_text)


def extract_text_from_docx(file_path: Path) -> str:
    """Extract all paragraph text from a DOCX file using python-docx."""
    logger.info("Extracting text from DOCX: %s", file_path)
    try:
        doc = Document(str(file_path))
    except Exception as exc:
        raise RuntimeError(f"Failed to read DOCX '{file_path}': {exc}") from exc

    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

    # Also pull text from tables (common in resume layouts)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and cell_text not in paragraphs:
                    paragraphs.append(cell_text)

    if not paragraphs:
        raise ValueError(f"No extractable text found in '{file_path}'.")

    logger.debug("  Extracted %d text blocks from DOCX", len(paragraphs))
    return "\n".join(paragraphs)


def extract_text(file_path: Path) -> str:
    """Dispatch to the correct extractor based on file extension."""
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(
            f"Unsupported file type '{suffix}'. Only PDF and DOCX are supported."
        )


# ---------------------------------------------------------------------------
# Gemini API call
# ---------------------------------------------------------------------------


def _clean_gemini_response(raw: str) -> str:
    """
    Robustly extract a JSON object from Gemini's response.

    Handles:
    - Bare JSON response
    - ```json ... ``` fenced blocks
    - Extra text / newlines before or after the JSON object
    - Partial JSON wrapped in markdown
    """
    # Step 1: strip markdown code fences
    text = re.sub(r"```(?:json)?", "", raw, flags=re.IGNORECASE)
    text = text.replace("```", "").strip()

    # Step 2: find the first { ... } block (greedy, handles nested braces)
    match = re.search(r"\{[\s\S]*\}", text)
    if match:
        return match.group().strip()

    # Step 3: return stripped text as-is and let json.loads raise a clear error
    return text.strip()


def _normalize_keys(obj):
    """Recursively clean dict keys (strip whitespace, quotes, newlines)."""
    if isinstance(obj, dict):
        new_dict = {}
        for k, v in obj.items():
            clean_k = str(k).strip(" \t\n\r\"'")
            new_dict[clean_k] = _normalize_keys(v)
        return new_dict
    elif isinstance(obj, list):
        return [_normalize_keys(item) for item in obj]
    return obj


def call_gemini(resume_text: str, api_key: str) -> dict:
    """
    Send the resume text to Gemini and return the parsed JSON dict.
    """
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(
    _GEMINI_MODEL,generation_config={"response_mime_type": "application/json"}
)

    prompt = EXTRACTION_PROMPT.format(resume_text=resume_text)
    logger.info("Sending resume text to Gemini (%d chars)…", len(resume_text))

    try:
        response = model.generate_content(prompt)
        raw_text = response.text
    except Exception as exc:
        raise RuntimeError(f"Gemini API call failed: {exc}") from exc

    logger.debug("Raw Gemini response:\n%s", raw_text)

    # Extract and parse JSON from Gemini's response
    cleaned = _clean_gemini_response(raw_text)
    try:
        data = json.loads(cleaned)
    except json.JSONDecodeError as exc:
        logger.error("JSON parsing error: %s. Raw output was: %r", exc, raw_text)
        raise RuntimeError(
            f"Failed to parse JSON from Gemini response: {exc}. "
            f"Raw snippet: {repr(raw_text[:200])}"
        ) from exc

    if not isinstance(data, dict):
        raise RuntimeError(
            f"Gemini returned JSON that is not an object/dict (got {type(data).__name__})"
        )

    return _normalize_keys(data)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def parse_resume(
    file_path: str | os.PathLike,
    api_key: Optional[str] = None,
) -> dict:
    """
    Parse a resume file (PDF or DOCX) into a structured JSON dict.

    Parameters
    ----------
    file_path : str | PathLike
        Absolute or relative path to the resume file.
    api_key : str, optional
        Google Generative AI API key. Falls back to the
        ``GEMINI_API_KEY`` environment variable if not provided.

    Returns
    -------
    dict
        Parsed resume with keys: name, email, phone, skills,
        experience, projects, education.

    Example
    -------
    >>> result = parse_resume("resume.pdf", api_key="YOUR_KEY")
    >>> print(result["name"])
    Jane Doe
    """
    path = Path(file_path).resolve()

    if not path.exists():
        raise FileNotFoundError(f"Resume file not found: '{path}'")

    # Resolve API key
    resolved_key = api_key or os.environ.get("GEMINI_API_KEY")
    if not resolved_key:
        raise EnvironmentError(
            "No Gemini API key provided. Pass api_key= or set the "
            "GEMINI_API_KEY environment variable."
        )

    # 1. Extract raw text
    raw_text = extract_text(path)
    logger.info("Extracted %d characters of text from '%s'", len(raw_text), path.name)

    # 2. Call Gemini to structure the data
    structured = call_gemini(raw_text, resolved_key)
    logger.info("Successfully parsed resume for: %s", structured.get("name", "Unknown"))

    return structured
